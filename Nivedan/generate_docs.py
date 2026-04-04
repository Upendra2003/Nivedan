"""
Nivedan — Technical Documentation Generator
Run: uv run --with python-docx python generate_docs.py
Output: Nivedan_Technical_Documentation.docx
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin   = Inches(1.2)
    section.right_margin  = Inches(1.2)

# ── Color palette ─────────────────────────────────────────────────────────────
BLUE       = RGBColor(0x1D, 0x4E, 0xD8)   # #1D4ED8
DARK_BLUE  = RGBColor(0x1E, 0x3A, 0x8A)   # #1E3A8A
TEAL       = RGBColor(0x0D, 0x9A, 0x88)   # #0D9A88
GRAY       = RGBColor(0x6B, 0x72, 0x80)   # #6B7280
LIGHT_GRAY = RGBColor(0xF3, 0xF4, 0xF6)
BLACK      = RGBColor(0x11, 0x18, 0x27)


# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Set background shading on a table cell."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)


def h1(text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(22)
    run.font.color.rgb = DARK_BLUE
    return p


def h2(text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(15)
    run.font.color.rgb = BLUE
    return p


def h3(text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(12)
    run.font.color.rgb = TEAL
    return p


def body(text: str, indent: float = 0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent   = Inches(indent)
    p.paragraph_format.space_after   = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = BLACK
    return p


def bullet(text: str, level: int = 0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent  = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    return p


def code_block(text: str):
    """Monospace shaded block for code snippets."""
    for line in text.strip().split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Inches(0.3)
        run = p.add_run(line if line else ' ')
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
        # light background via shading
        pPr  = p._p.get_or_add_pPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  'F1F5F9')
        pPr.append(shd)
    doc.add_paragraph()   # blank line after block


def divider():
    p = doc.add_paragraph('─' * 90)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0xCB, 0xD5, 0xE1)
        run.font.size      = Pt(7)


def two_col_table(rows: list[tuple[str, str]], header: tuple[str, str] = None):
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    table.autofit = False
    table.columns[0].width = Inches(2.2)
    table.columns[1].width = Inches(3.8)
    if header:
        hrow = table.add_row()
        for i, h in enumerate(header):
            cell = hrow.cells[i]
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            set_cell_bg(cell, '1D4ED8')
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for col_a, col_b in rows:
        row = table.add_row()
        row.cells[0].text = col_a
        row.cells[1].text = col_b
        row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
        row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)
    doc.add_paragraph()


# ═══════════════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run('Nivedan')
r.bold = True
r.font.size = Pt(40)
r.font.color.rgb = DARK_BLUE

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub_p.add_run('Technical Documentation')
r2.font.size = Pt(20)
r2.font.color.rgb = BLUE

doc.add_paragraph()
tag_p = doc.add_paragraph()
tag_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = tag_p.add_run('AI-Powered Civic Complaint Filing for Indian Citizens')
r3.font.size = Pt(13)
r3.italic = True
r3.font.color.rgb = GRAY

doc.add_paragraph()
doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
mr = meta.add_run('Stack: Flask · React Native (Expo) · React · MongoDB · Sarvam AI')
mr.font.size = Pt(10)
mr.font.color.rgb = GRAY

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 1. SYSTEM OVERVIEW
# ═══════════════════════════════════════════════════════════════════════════════

h1('1. System Overview')
body(
    'Nivedan is a multilingual AI assistant that guides Indian citizens through the process '
    'of filing government complaints — without requiring legal literacy or fluency in English. '
    'A user describes their problem in their mother tongue, and the AI agent gathers all required '
    'information conversationally, fills the correct government form as a PDF, collects a '
    'handwritten signature and supporting documents, then submits the complete package to the '
    'relevant authority.'
)

h2('1.1 User Journey')
steps = [
    ('Step 1 — Category Selection', 'User opens the mobile app and picks a legal category (e.g. "Labor Issues > Salary Not Paid") from a grid.'),
    ('Step 2 — AI Conversation',    'An AI agent (powered by Sarvam-M) greets the user in their preferred language and understands the problem through natural dialogue.'),
    ('Step 3 — Blank Form Preview', 'The agent fetches and shows the blank government form so the user knows what will be filed.'),
    ('Step 4 — Field Collection',   'The agent collects required fields conversationally — name, employer, dates, amounts — auto-filling what it already knows from the profile.'),
    ('Step 5 — Signature & Docs',   'User uploads a photo of their handwritten signature and any supporting documents (pay slips, bank statements, etc.).'),
    ('Step 6 — PDF Preview',        'A professional filled PDF is generated server-side and shown in-chat. User approves or requests changes.'),
    ('Step 7 — Submission',         'The complete PDF (with embedded signature and attached documents) is submitted to the mock government portal.'),
    ('Step 8 — Status Updates',     'The portal processes the case and fires webhooks; the user receives Expo push notifications and in-app status cards.'),
]
two_col_table(steps, header=('Step', 'Description'))

h2('1.2 Supported Languages')
body('English, Hindi (हिन्दी), Tamil (தமிழ்), Telugu (తెలుగు), Kannada (ಕನ್ನಡ), Malayalam (മലയാളം), Bengali, Marathi, Gujarati, Punjabi — via Sarvam AI model sarvam-m.')

h2('1.3 Supported Complaint Categories')
cats = [
    ('salary_not_paid',       'Salary Non-Payment → Labour Commissioner'),
    ('wrongful_termination',  'Wrongful Termination → Labour Commissioner'),
    ('workplace_harassment',  'Workplace Harassment → Labour Commissioner / ICC'),
    ('fir_not_registered',    'FIR Not Registered → Superintendent of Police'),
    ('police_misconduct',     'Police Misconduct → SP / State Human Rights Commission'),
    ('defective_product',     'Defective Product → Consumer Disputes Redressal Forum'),
    ('online_scam',           'Online Scam / Cyber Fraud → Cyber Crime Cell'),
    ('(any other)',           'Generic complaint → Relevant Authority (fallback config)'),
]
two_col_table(cats, header=('Subcategory ID', 'Title → Authority'))

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 2. HIGH-LEVEL ARCHITECTURE
# ═══════════════════════════════════════════════════════════════════════════════

h1('2. High-Level Architecture')
body(
    'Nivedan is composed of four independently runnable processes that communicate '
    'over HTTP. MongoDB is the single shared data store.'
)

arch_rows = [
    ('Mobile App',       'React Native + Expo SDK 54',  'Expo Go / dev build', 'User-facing: chat, PDF preview, notifications'),
    ('Backend API',      'Flask 3 (Python)',             'http://0.0.0.0:5000', 'Auth, agent, complaints, PDF generation, webhooks'),
    ('Mock Gov Portal',  'Flask 3 (Python)',             'http://localhost:5001','Simulates government e-filing portal, fires webhooks'),
    ('Web Dashboard',    'React 18 + Vite',              'http://localhost:5173','Staff / admin view of cases, notifications, analytics'),
]
tbl = doc.add_table(rows=0, cols=4)
tbl.style = 'Table Grid'
hrow = tbl.add_row()
for i, h in enumerate(['Component', 'Framework', 'Port / URL', 'Responsibility']):
    c = hrow.cells[i]
    c.text = h
    c.paragraphs[0].runs[0].bold = True
    c.paragraphs[0].runs[0].font.size = Pt(10)
    set_cell_bg(c, '1E3A8A')
    c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
for row_data in arch_rows:
    row = tbl.add_row()
    for i, val in enumerate(row_data):
        row.cells[i].text = val
        row.cells[i].paragraphs[0].runs[0].font.size = Pt(10)
doc.add_paragraph()

h2('2.1 Data Flow')
body('Below is the request/response flow for a complete complaint filing:')
code_block('''\
Mobile App
  │  POST /complaints/create          → Backend creates Mongo doc
  │  POST /agent/message {msg:""}     → Backend returns greeting
  │  POST /agent/message {msg:"..."}  → LLM conversation (Sarvam AI)
  │  POST /agent/message              → Backend generates PDF (fpdf2)
  │  POST /complaints/<id>/upload-doc → Signature / supporting doc stored in Mongo
  │  POST /agent/message {msg:"yes"}  → Backend submits to Mock Portal
  │
Mock Portal
  │  POST /webhooks/portal            → Backend receives status update
  │
Backend
  │  POST Expo Push API               → Mobile receives push notification
  │
Mobile App  ←─ polling fallback (GET /complaints/<id> every 10s)
''')

h2('2.2 Database')
body('MongoDB (local). Two databases:')
two_col_table([
    ('Nivedan',        'users, complaints, notifications — used by Backend API'),
    ('Nivedan_portal', 'portal_submissions — used by Mock Portal'),
], header=('Database', 'Purpose'))

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 3. BACKEND — Flask API
# ═══════════════════════════════════════════════════════════════════════════════

h1('3. Backend — Flask API')
body(
    'Located at Nivedan/backend/. Managed exclusively with uv (never pip). '
    'Entry point: app.py. All routes registered as Flask blueprints. '
    'MongoDB accessed via raw pymongo — no ORM.'
)

h2('3.1 Entry Point — app.py')
body(
    'Loads .env via python-dotenv before any other import (critical so db.py and '
    'routes read env vars at import time). Fails loudly if JWT_SECRET is missing. '
    'Warns at startup if WEBHOOK_SECRET is unset. Enables CORS for all origins '
    '(dev mode). Binds Flask to host 0.0.0.0 so mobile devices on the LAN can reach it.'
)
code_block('''\
# app.py (simplified)
import os
from dotenv import load_dotenv
load_dotenv()                              # MUST be first — before any import

if not os.getenv("JWT_SECRET"):
    raise RuntimeError("JWT_SECRET not set. Refusing to start.")

if not os.getenv("WEBHOOK_SECRET"):
    import warnings
    warnings.warn("WEBHOOK_SECRET not set — webhooks accept unauthenticated requests.")

from flask import Flask
from flask_cors import CORS
from routes.auth      import auth_bp
from routes.complaints import complaints_bp
# ... (agent_bp, forms_bp, notifications_bp, users_bp, webhooks_bp)

app = Flask(__name__)
CORS(app)
app.register_blueprint(auth_bp,        url_prefix="/auth")
app.register_blueprint(complaints_bp,  url_prefix="/complaints")
# ... register remaining blueprints

if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000,
            debug=os.getenv("FLASK_DEBUG","false").lower()=="true")
''')

h2('3.2 Routes')

h3('routes/auth.py')
body('Handles user registration, login, and token verification.')
two_col_table([
    ('POST /auth/register', 'Create user. Hashes password with bcrypt. Returns JWT token.'),
    ('POST /auth/login',    'Verify email + password. Returns JWT token + user profile.'),
    ('GET  /auth/me',       'Decode JWT → return { id, name, email, preferred_language }.'),
], header=('Endpoint', 'Description'))
code_block('''\
# Registration stores: name, email, phone, bcrypt(password), preferred_language
# JWT payload: { user_id: str(ObjectId), exp: 30 days }
# All tokens use HS256 with JWT_SECRET from .env
''')

h3('routes/complaints.py')
body('Full complaint lifecycle: create, list, retrieve, status updates, portal submission, and document upload.')
two_col_table([
    ('POST /complaints/create',              'Create complaint doc; initialise agent_state=CHAT, agent_history=[].'),
    ('GET  /complaints/',                    'List all complaints for the authenticated user.'),
    ('GET  /complaints/mine',               'List + counts (pending/filed/resolved) for dashboard.'),
    ('GET  /complaints/<id>',               'Full doc including timeline, form_data, agent_history.'),
    ('POST /complaints/<id>/submit_to_portal', 'Send pdf_base64 to Mock Portal, record portal_ref_id.'),
    ('POST /complaints/<id>/resubmit',      'Re-send after rejection; increments resubmission_count.'),
    ('PATCH /complaints/<id>/status',       'Internal/webhook: update status + append to timeline.'),
    ('POST /complaints/<id>/upload-doc',    'Store signature or supporting document in MongoDB.'),
], header=('Endpoint', 'Description'))

body('Document upload endpoint (Phase 9):')
code_block('''\
# POST /complaints/<id>/upload-doc
# Body: { type: "signature"|"supporting", file_base64, filename, mime_type }
#
# Server-side guards:
_ALLOWED_MIME_TYPES = {
    "image/jpeg","image/jpg","image/png","image/gif",
    "image/webp","application/pdf"
}
_MAX_SUPPORTING_DOCS = 10

# Signature → stored in complaint.signature_b64
# Supporting doc → $push into complaint.supporting_docs[]
# Rejects: wrong MIME (415), too many docs (400), file > 1.2 MB base64 (413)
''')

h3('routes/agent.py')
body('Single endpoint that drives the entire AI conversation loop.')
code_block('''\
# POST /agent/message
# Body:    { complaint_id: str, message: str }
# Returns: {
#   reply:         str,           # AI response text (stripped of <think> tags)
#   action:        str | null,    # "show_pdf" | "show_buttons" | "status_update" | null
#   action_data:   {              # present when action != null
#     filename?:    str,
#     pdf_base64?:  str,
#     label?:       str,          # "blank_form" or "filled_form"
#     buttons?:     str[],
#     status?:      str,
#     portal_ref_id?: str,
#   },
#   thinking_steps: str[],        # displayed by ThinkingStrip component
# }
''')

h3('routes/forms.py')
body('PDF serving and generation.')
two_col_table([
    ('GET  /forms/<form_name>', 'Proxy to Mock Portal — returns blank form PDF bytes.'),
    ('POST /forms/generate',   'Generate filled PDF: { form_name, form_data } → { pdf_base64 }.'),
], header=('Endpoint', 'Description'))

h3('routes/notifications.py')
two_col_table([
    ('GET  /notifications/mine',         'Unread notifications. Add ?all=1 to return all (read + unread).'),
    ('POST /notifications/read/<id>',    'Mark single notification as read.'),
    ('POST /notifications/read/all',     'Mark all notifications as read.'),
], header=('Endpoint', 'Description'))

h3('routes/users.py')
body('POST /users/push_token — Stores Expo push token on the user document for push notifications.')

h3('routes/webhooks.py')
body('POST /webhooks/portal — Called by Mock Portal when case status changes. No JWT auth (uses WEBHOOK_SECRET). Updates complaint status, appends timeline, creates notification, triggers Expo push.')

h2('3.3 Services')

h3('services/auth_middleware.py')
body('@jwt_required decorator — decodes JWT, loads full user doc from MongoDB, sets flask.g.user. Applied to all authenticated routes.')
code_block('''\
@jwt_required
def my_route():
    user = g.user          # full user dict from MongoDB
    user_id = user["_id"]  # ObjectId
''')

h3('services/sarvam.py')
body('Wraps the sarvamai SDK for LLM completions. Only chat_completion() is implemented; audio transcription/TTS stubs were removed during refactoring.')
code_block('''\
def chat_completion(messages: list, system_prompt: str = "") -> str:
    client = SarvamAI(api_subscription_key=SARVAM_API_KEY)
    all_msgs = [{"role":"system","content":system_prompt}] + messages
    response = client.chat.completions(
        messages=all_msgs,
        model="sarvam-m",
        reasoning_effort=None,   # suppress CoT at model level
    )
    return response.choices[0].message.content or ""
''')

h3('services/agent_runner.py')
body(
    'The heart of Nivedan. Implements a goal-oriented state machine that drives '
    'the entire complaint filing flow. Each call to run_agent() loads the complaint '
    'from MongoDB, determines the current state, calls the LLM, parses the output, '
    'and persists state back to MongoDB.'
)
body('State machine transitions:')
code_block('''\
CHAT
  │  LLM outputs ACTION: fetch_form
  ▼
SHOW_BLANK_FORM
  │  User says "fill it" → LLM outputs ACTION: collect_fields
  ▼
COLLECT_FIELDS
  │  All required fields gathered → Python triggers ACTION: fill_form
  ▼
COLLECT_DOCS      ← NEW (Phase 9)
  │  Sub-stage 1: signature  → user uploads or skips
  │  Sub-stage 2: documents  → user uploads files, clicks "Done"
  │  User says "done" → generate full PDF
  ▼
PREVIEW
  │  User confirms → LLM outputs ACTION: submit
  ▼
SUBMITTED
''')

body('LLM output format (every response must end with):')
code_block('''\
EXTRACTED: {"field_key": "value_from_conversation"}
ACTION: none|fetch_form|collect_fields|fill_form|submit
''')

body('Key implementation details:')
bullet('_strip_think() removes <think>...</think> blocks (sarvam-m emits chain-of-thought). Handles both closed tags and unclosed <think> tags.')
bullet('Auto-prefills complainant_name from user.name on first field collection.')
bullet('SUBCATEGORY_CONFIG dict — adding a new complaint type only needs a new dict entry.')
bullet('State machine guards prevent LLM from skipping states (e.g. "submit" from COLLECT_FIELDS → forced to fill_form first).')
bullet('agent_history capped at 40 messages; leading assistant messages stripped before every LLM call (Sarvam requires user-first).')

h3('services/pdf_generator.py')
body('Server-side PDF generation using fpdf2. Never renders on client.')
two_col_table([
    ('generate_salary_complaint(data, signature_b64, supporting_docs)', 'Professional A4 form: 6 sections, embedded signature (EXIF-corrected), auto-ticked document checkboxes, supporting docs appended as pages.'),
    ('_orient_image(img_bytes)',       'Applies EXIF orientation using Pillow ImageOps.exif_transpose() — fixes phone photos that are 90° rotated.'),
    ('_detect_doc_type(filename)',     'Maps filename to Section-5 checkbox key (salary_slip, bank_statement, etc.).'),
    ('_append_signature_and_docs(pdf_bytes, signature_b64, supporting_docs)', 'Generic helper for non-salary forms — appends signature page + documents via pypdf merge.'),
    ('generate_pdf_b64(form_name, data, signature_b64, supporting_docs)', 'Entry point from agent_runner. Routes to salary-specific or generic generator.'),
], header=('Function', 'Purpose'))

code_block('''\
# Salary complaint form — Section 5 auto-tick based on uploaded filenames
doc_types = {_detect_doc_type(d.get("filename","")) for d in (supporting_docs or [])}
checkbox("Salary Slips / Pay Stubs", "salary_slip" in doc_types)
checkbox("Bank Statements",          "bank_statement" in doc_types)

# Signature embedding (EXIF-corrected, size-guarded)
if signature_b64:
    decoded = base64.b64decode(signature_b64)
    if len(decoded) <= 1_200_000:          # MED-3: cap Pillow memory usage
        raw = _orient_image(decoded)
        pdf.image(BytesIO(raw), x=65, y=sig_y-1, w=48, h=10)
''')

h3('services/form_handler.py')
body('Phase 1b legacy helper. Provides FIELD_DEFINITIONS, get_fields, next_missing_field, build_summary, submit_to_portal. Still used by the old /complaints/<id>/chat route.')

h2('3.4 Models')
body('Schema factory functions in models/. Raw dicts — no ORM. Passed to db.collection.insert_one().')
code_block('''\
# models/complaint.py
def complaint_schema(user_id, category, subcategory, form_data, title):
    return {
        "user_id":            user_id,
        "category":           category,
        "subcategory":        subcategory,
        "title":              title,
        "status":             "pending",
        "form_data":          form_data,
        "agent_state":        "CHAT",
        "agent_history":      [],
        "signature_b64":      None,    # stored after upload
        "supporting_docs":    [],      # list of {filename, file_b64, mime_type}
        "portal_ref_id":      None,
        "rejection_reason":   None,
        "resubmission_count": 0,
        "current_step_label": "",
        "timeline":           [],
        "documents":          [],
        "created_at":         datetime.now(timezone.utc),
        "updated_at":         datetime.now(timezone.utc),
    }
''')

h2('3.5 db.py')
body('Single MongoDB connection. Import db from here in every route/service.')
code_block('''\
from pymongo import MongoClient
client = MongoClient(os.getenv("MONGO_URI", "mongodb://localhost:27017/Nivedan"))
db     = client.get_default_database()
# Usage: db.complaints.find_one({"_id": ObjectId(id)})
''')

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 4. MOBILE APP — React Native / Expo
# ═══════════════════════════════════════════════════════════════════════════════

h1('4. Mobile App — React Native / Expo')
body(
    'Located at Nivedan/mobile/. Expo SDK 54 with expo-router file-based routing. '
    'Runs in Expo Go for development; a dev build is required for react-native-pdf (PDF viewer). '
    'API base URL read from EXPO_PUBLIC_API_URL in mobile/.env.'
)

h2('4.1 Navigation Structure')
code_block('''\
app/
  _layout.tsx             Root: SafeAreaProvider > AuthProvider > NotificationProvider > Stack
                          Auth guard: unauthenticated → /auth/login
  auth/
    login.tsx             Email + password login
    register.tsx          Register with language chip picker (6 languages)
  (tabs)/
    _layout.tsx           Bottom tab bar: Home / My Cases / Profile
    index.tsx             Home: category grid + subcategory bottom sheet + recent cases
    dashboard.tsx         My Cases: complaint list with status badges
    profile.tsx           User info + language + sign out
  chat/
    [category].tsx        AI agent chat (main screen — see §4.3)
  pdf-viewer.tsx          Full-screen PDF viewer (native only, react-native-pdf)
''')

h2('4.2 Screens')

h3('app/(tabs)/index.tsx — Home Screen')
body('Category grid (6 complaint categories). Tapping a category opens a bottom sheet (Modal + Animated.spring) listing subcategories. Tapping a subcategory navigates to /chat/<subcategory_id>. Also shows recent cases from GET /complaints/mine.')
code_block('''\
// Navigate to chat with subcategory ID
router.push(`/chat/${subcategoryId}`)
// e.g. /chat/salary_not_paid
''')

h3('app/chat/[category].tsx — AI Chat Screen')
body(
    'The most complex screen. Manages the entire agent conversation: '
    'creates the complaint, sends the greeting, handles all message types, '
    'uploads signatures and documents, renders the PDF preview, and polls for status changes.'
)
body('Key state variables:')
code_block('''\
const [messages,       setMessages]       = useState<Message[]>([])
const [stage,          setStage]          = useState<Stage>("loading")
const [thinking,       setThinking]       = useState(false)   // AI is processing
const [uploading,      setUploading]      = useState(false)   // file uploading
const [thinkingSteps,  setThinkingSteps]  = useState<string[]>()
const [complaintId,    setComplaintId]    = useState<string | null>(null)
const [input,          setInput]          = useState("")
''')
body('Initialization flow:')
code_block('''\
useEffect(() => {
  // 1. Create complaint doc
  const doc = await api.authedPost("/complaints/create", { category, subcategory, form_data:{} })
  setComplaintId(doc._id)   // NOTE: _id NOT complaint_id

  // 2. Send empty message → get greeting (no LLM call, hardcoded in agent_runner)
  const res = await sendAgentMessage(doc._id, "")
  applyResponse(res)
}, [])
''')
body('Document upload handlers:')
code_block('''\
// Signature — opens gallery, uploads to /complaints/<id>/upload-doc
const handleSignatureUpload = async () => {
  const result = await ImagePicker.launchImageLibraryAsync({ base64: true, quality: 0.5 })
  if (result.canceled) return
  const b64 = result.assets[0].base64!
  if (b64.length > MAX_B64) { /* show error */ return }
  setUploading(true)   // ← uploading state, NOT thinking — different UX
  await api.authedPost(`/complaints/${complaintId}/upload-doc`, {
    type: "signature", file_base64: b64, filename: "signature.jpg", mime_type: "image/jpeg"
  })
  pushMsg({ id: uid(), type: "uploaded_file_card", filename: "signature.jpg", isSignature: true })
  setUploading(false)
  await handleSend("signature uploaded")  // this triggers thinking=true for AI
}

// Document — camera / gallery / file picker
const handleDocumentUpload = async (source: "camera" | "gallery" | "file") => {
  // Open picker BEFORE setUploading — user is browsing, not waiting on us
  const result = await ImagePicker.launchImageLibraryAsync({ base64: true, quality: 0.4 })
  if (result.canceled) return
  setUploading(true)
  await api.authedPost(`/complaints/${complaintId}/upload-doc`, { type: "supporting", ... })
  // ...
}
''')

h3('app/pdf-viewer.tsx — Full-Screen PDF Viewer')
body('Native-only route. Reads PDF data from utils/pdfStore.ts (module-level store). Writes base64 → temp file via expo-file-system. Renders with react-native-pdf. Bottom bar has "Request Changes" and "Approve & Submit" buttons.')

h2('4.3 Components')

h3('components/ChatMessageRow.tsx')
body('Renders a single message bubble. Pure switch on msg.type. All 8 message types:')
two_col_table([
    ('user',                    'Right-aligned blue bubble with user text.'),
    ('agent',                   'Left-aligned bubble with "Nivedan AI" label + dot.'),
    ('action_buttons',          'Row of pill buttons; "yes/submit/confirm" buttons get primary color.'),
    ('status_update',           'Green card with tick icon — "Complaint Filed!"'),
    ('pdf_card',                'Document card with View PDF + Submit buttons. blank_form hides Submit.'),
    ('signature_request',       'Card with Upload / Skip buttons. Inline uploading state while file uploads.'),
    ('document_upload_request', 'Card with Camera / Gallery / File buttons + Done button.'),
    ('uploaded_file_card',      'Small row showing uploaded filename with green checkmark.'),
], header=('msg.type', 'Renders'))

h3('components/ThinkingStrip.tsx')
body('Claude-like thinking display. Shown while AI is processing. Cycles through "Thinking.../Processing..." with pulsing dot. After AI responds, animates thinking steps in one-by-one.')

h3('components/InAppBanner.tsx')
body('Slides down from top when a push notification arrives while the app is open. Auto-dismisses after 4 seconds. Spring animation.')

h3('components/NotificationDrawer.tsx')
body('Right-side slide-in drawer showing all notifications. Unread items highlighted. Shows time-ago, type icon, and mark-read functionality.')

h2('4.4 Services')

h3('services/api.ts')
body('All HTTP calls go through api.get/post/authedGet/authedPost/authedPatch. Automatically reads EXPO_PUBLIC_API_URL → Metro hostUri → localhost:5000.')

h3('services/agent.ts')
body('sendAgentMessage(complaintId, message) → AgentResponse. Wraps POST /agent/message.')

h3('services/notifications.ts')
body('registerForPushNotifications(), addForegroundListener(), addResponseListener(). Handles Expo notification permissions and token registration.')

h2('4.5 Hooks')

h3('hooks/useComplaintPolling.ts')
body(
    'Two polling effects. (1) Rejection poll — every 8 s when stage="submitted", '
    'detects REJECTED state and sends empty agent message to trigger resubmission. '
    '(2) Status fallback — every 10 s, inserts status_update card if status changed. '
    'Both use a cancelled flag to prevent state updates after unmount.'
)

h2('4.6 Context & Utilities')
two_col_table([
    ('context/AuthContext.tsx',       'AuthProvider + useAuth() — login/register/logout + session restore from secure storage.'),
    ('context/NotificationContext.tsx', 'NotificationProvider + useNotifications() — polls /notifications/mine every 30 s, provides unreadCount, markRead, markAllRead.'),
    ('utils/pdfStore.ts',             'Module-level store for PDF data between screens (avoids URL-param size limits).'),
    ('utils/storage.ts',              'Platform-aware wrapper: web → localStorage, native → expo-secure-store.'),
    ('constants/categories.ts',       'CATEGORIES array + findSubcategory(subcategoryId) helper.'),
    ('constants/theme.ts',            'darkTheme, lightTheme, Theme type, useTheme() hook (reads system color scheme).'),
], header=('File', 'Purpose'))

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 5. WEB DASHBOARD — React / Vite
# ═══════════════════════════════════════════════════════════════════════════════

h1('5. Web Dashboard — React / Vite')
body(
    'Located at Nivedan/web/. React 18 + react-router-dom + Tailwind CSS (CDN) + Recharts. '
    'Target: staff / admin monitoring. Dark mode by default (darkMode: "class" on <html>).'
)

h2('5.1 Pages')

h3('src/pages/Login.tsx & Register.tsx')
body('Login with floating dark/light theme toggle. Register includes a language chip picker (6 language options shown as selectable pills).')

h3('src/pages/Dashboard.tsx')
body('4 stat cards (total / pending / resolved / failed). Donut chart (Recharts PieChart) showing status distribution. Complaint list with ProgressBar per case. Recent notification feed.')

h3('src/pages/CaseDetail.tsx')
body('Full case view: timeline (latest event first), rejection panel with resubmit option, uploaded documents list, form_data table, PDF download link.')

h3('src/pages/Notifications.tsx')
body('Filter tabs (All / Unread / Status Updates / Rejections). Mark-all-read button. Each notification card links to the relevant case.')

h2('5.2 Layout & Navigation')
body('src/components/Layout.tsx — sidebar navigation + top bar (logo, theme toggle, avatar, sign-out). Fetches /notifications/mine for badge count on sidebar bell icon.')

h2('5.3 Auth & API')
code_block('''\
// src/services/api.ts
// In-memory token (_token module variable — NOT localStorage for security)
let _token: string | null = null

export const setToken = (t: string) => { _token = t }

async function request(path: string, opts: RequestInit = {}) {
  const res = await fetch(API_BASE + path, {
    ...opts,
    headers: {
      "Content-Type": "application/json",
      ...(opts.headers || {}),
      ...(_token ? { Authorization: `Bearer ${_token}` } : {}),
    },
  })
  if (!res.ok) throw new Error(await res.text())
  return res.json()
}
''')

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 6. MOCK GOVERNMENT PORTAL
# ═══════════════════════════════════════════════════════════════════════════════

h1('6. Mock Government Portal')
body(
    'Located at Nivedan/mock_portal/portal_app.py. Simulates a government e-filing portal. '
    'Separate Flask process on port 5001. Has its own MongoDB database (Nivedan_portal). '
    'Dark-theme dashboard with case table, step dots, and action modals. Auto-refreshes every 15 s.'
)

h2('6.1 Endpoints')
two_col_table([
    ('GET  /portal/dashboard',        'Renders the HTML dashboard (templates/dashboard.html).'),
    ('POST /portal/submit',           'Receives complaint + PDF from Backend. Creates portal_submission doc. Returns { portal_ref_id }.'),
    ('GET  /portal/submissions',      'JSON list of all portal submissions (for API consumers).'),
    ('POST /portal/action',           'Admin approves / advances / fails a case. Fires webhook back to Backend.'),
    ('GET  /portal/pdf/<ref_id>',     'Download the submitted PDF for a case.'),
    ('GET  /portal/forms/salary_complaint', 'Serves the blank salary complaint form PDF (generated with fpdf2).'),
], header=('Endpoint', 'Description'))

h2('6.2 Webhook Flow')
body('When an admin clicks Approve/Advance/Fail in the dashboard, portal fires:')
code_block('''\
POST http://backend:5000/webhooks/portal
{
  "portal_ref_id": "CF-2024-001",
  "status":        "resolved",      # or "under_review", "next_step", "failed"
  "next_step":     "Payment to be processed within 30 days",
  "reason":        ""               # populated on failure/rejection
}
''')
body('Backend webhooks.py processes this: updates complaint status, appends timeline entry, creates notification, sends Expo push notification.')

h2('6.3 Status Progression')
body('Cases move through: pending → filed → acknowledged → under_review → next_step → resolved / failed.')

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 7. AI AGENT SYSTEM
# ═══════════════════════════════════════════════════════════════════════════════

h1('7. AI Agent System')
body('The agent is a Python state machine in services/agent_runner.py, driven by the Sarvam-M LLM. Each user message goes through run_agent() which reads complaint state from MongoDB, makes an LLM call if needed, parses the output, and writes updated state back.')

h2('7.1 State Machine')
code_block('''\
State: CHAT
  Purpose: Understand the user\'s problem through natural conversation.
  Trigger to advance: LLM outputs "ACTION: fetch_form"
  What happens: Agent fetches blank PDF from portal, stores URL.

State: SHOW_BLANK_FORM
  Purpose: Show user the blank government form.
  Trigger to advance: User says "fill it" / "yes" / "proceed"
  What happens: LLM outputs "ACTION: collect_fields"

State: COLLECT_FIELDS
  Purpose: Collect all required fields conversationally.
  How: LLM extracts values via "EXTRACTED: {}" at end of each response.
  Trigger to advance: All required fields present in form_data.
  What happens: Python generates PDF, transitions to COLLECT_DOCS.

State: COLLECT_DOCS
  Sub-stage "signature":
    Shows signature_request card.
    Advances when: complaint.signature_b64 is set (upload) OR user says "skip".
  Sub-stage "documents":
    Shows document_upload_request card.
    Advances when: user says "done" / "finish" / "generate" / "proceed".
  What happens on advance: Regenerates PDF with signature + docs, shows in PREVIEW.

State: PREVIEW
  Purpose: User reviews filled PDF.
  Trigger to advance: User confirms ("yes, submit it").
  What happens: LLM outputs "ACTION: submit"

State: SUBMITTED
  Purpose: Calls portal, marks complaint filed.
  Returns: action="status_update" with portal_ref_id.
''')

h2('7.2 Subcategory Configuration')
body('SUBCATEGORY_CONFIG in agent_runner.py drives everything: greeting, system prompt, fields to collect, form filename, and authority name. Adding a new complaint type requires only a new dict entry:')
code_block('''\
SUBCATEGORY_CONFIG = {
    "salary_not_paid": {
        "title":     "Salary Non-Payment",
        "authority": "Labour Commissioner",
        "issue":     "salary not being paid",
        "form_name": "salary_non_payment",
        "fields": [
            ("complainant_name",    "your full name"),
            ("employer_name",       "your employer / company name"),
            ("months_pending",      "number of months salary is unpaid"),
            ("amount_pending",      "total unpaid amount in rupees"),
            # ... more fields
        ],
    },
    # Add new category here — everything else is automatic
}
''')

h2('7.3 Think-Tag Stripping')
body('Sarvam-M emits chain-of-thought inside <think>...</think> tags. These must be stripped before sending to the client. The _strip_think() function handles both closed and unclosed tags:')
code_block('''\
def _strip_think(text: str) -> str:
    if "</think>" in text:
        # Take everything after the last closing tag
        return text.split("</think>")[-1].strip()
    # Unclosed <think> — remove the tag itself but keep remaining text
    return re.sub(r"<think>", "", text).strip()
''')

h2('7.4 Field Extraction')
body('After every LLM response in COLLECT_FIELDS, the agent parses the EXTRACTED block:')
code_block('''\
# LLM response example:
# "I see, you have been working at Infosys since January 2023.
#  EXTRACTED: {"employer_name": "Infosys", "employment_start_date": "January 2023"}
#  ACTION: none"

extracted = json.loads(extracted_json)
for key, val in extracted.items():
    if val and key in valid_keys and len(str(val)) < 500:
        form_data[key] = val

db.complaints.update_one(
    {"_id": ObjectId(complaint_id)},
    {"$set": {"form_data": form_data}}
)
''')

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 8. DOCUMENT UPLOAD & PDF GENERATION
# ═══════════════════════════════════════════════════════════════════════════════

h1('8. Document Upload & PDF Generation')

h2('8.1 Upload Flow')
code_block('''\
Mobile                       Backend                         MongoDB
  │                            │                               │
  │─ POST /upload-doc ────────►│                               │
  │  { type:"signature",       │── validate MIME type ────────►│
  │    file_base64,            │   validate size (≤1.2MB b64)  │
  │    filename,               │   validate doc count (≤10)    │
  │    mime_type }             │                               │
  │                            │── $set signature_b64 ────────►│
  │◄─ { ok: true } ───────────│   (or $push supporting_docs)  │
  │                            │                               │
  │─ "done uploading" ────────►│                               │
  │  (agent message)           │── load complaint ────────────►│
  │                            │   generate_pdf_b64(           │
  │                            │     form_name, form_data,     │
  │                            │     signature_b64,            │
  │                            │     supporting_docs           │
  │                            │   )                           │
  │◄─ { action:"show_pdf" } ──│                               │
''')

h2('8.2 PDF Generation Pipeline')
code_block('''\
generate_pdf_b64(form_name, form_data, signature_b64, supporting_docs)
  │
  ├─ "salary" in form_name → generate_salary_complaint(...)
  │     │
  │     ├─ Build 6-section A4 form (fpdf2)
  │     ├─ Section 5: auto-tick checkboxes from filename heuristics
  │     │   _detect_doc_type("salary_slip_march.jpg") → "salary_slip"
  │     │   _detect_doc_type("bank_statement.pdf")    → "bank_statement"
  │     │
  │     ├─ Embed signature at declaration line (EXIF-corrected)
  │     │   decoded = base64.b64decode(signature_b64)
  │     │   if len(decoded) <= 1_200_000:      # size guard
  │     │       raw = _orient_image(decoded)   # fix phone rotation
  │     │       pdf.image(BytesIO(raw), x=65, y=sig_y-1, w=48, h=10)
  │     │
  │     ├─ Append image docs as new fpdf2 pages (> 700 KB → placeholder)
  │     └─ Merge PDF docs using pypdf PdfWriter
  │
  └─ Other forms → generate_pdf(generic fallback)
       └─ _append_signature_and_docs(pdf_bytes, signature_b64, supporting_docs)
            └─ Same signature + doc append logic via pypdf merge
''')

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 9. NOTIFICATION SYSTEM
# ═══════════════════════════════════════════════════════════════════════════════

h1('9. Notification System (Phase 8)')
body('Two-layer notification system: Expo push notifications (primary) + 10-second polling fallback (for demo environments where push may not arrive).')

h2('9.1 Push Notification Flow')
code_block('''\
1. Mobile registers → POST /users/push_token  (Expo push token stored on user doc)
2. Portal fires POST /webhooks/portal
3. Backend webhooks.py:
   a. Updates complaint status + timeline
   b. Creates notification doc in MongoDB
   c. Sends Expo push via _send_expo_push():
      POST https://exp.host/--/api/v2/push/send
      {
        to:    user.push_token,
        title: "Nivedan Update",
        body:  "Your salary complaint is now Under Review",
        sound: "default",
        data:  { type, complaint_id, status, next_step_label }
      }
''')

h2('9.2 Mobile Notification Handling')
code_block('''\
// context/NotificationContext.tsx
// Provides: notifications, unreadCount, refreshNotifications, markRead, markAllRead
// Polls: GET /notifications/mine?all=1  every 30 seconds

// services/notifications.ts
registerForPushNotifications()   // request permission + POST /users/push_token
addForegroundListener(cb)        // push arrives while app is open → InAppBanner
addResponseListener(cb)          // user taps a push → navigate to case

// Notification handler set to suppress system banner (app shows InAppBanner instead)
Notifications.setNotificationHandler({
  handleNotification: async () => ({
    shouldShowAlert: false,   // we show InAppBanner instead
    shouldPlaySound: true,
    shouldSetBadge:  true,
  }),
})
''')

h2('9.3 Polling Fallback (useComplaintPolling)')
code_block('''\
// Poll 1: Rejection detection (every 8 s when stage="submitted")
setInterval(async () => {
  const c = await api.authedGet(`/complaints/${complaintId}`)
  if (c.agent_state === "REJECTED") {
    // Send empty agent message → agent auto-corrects form and regenerates PDF
    const res = await sendAgentMessage(complaintId, "")
    applyResponse(res)
  }
}, 8000)

// Poll 2: Status fallback (every 10 s at all non-loading stages)
setInterval(async () => {
  const c = await api.authedGet(`/complaints/${complaintId}`)
  if (prevStatus !== c.status) {
    pushStatusCard(c.status, c.current_step_label)
    refreshNotifications()
  }
}, 10_000)
''')

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 10. DATABASE SCHEMAS
# ═══════════════════════════════════════════════════════════════════════════════

h1('10. Database Schemas')

h2('10.1 users collection (Nivedan)')
code_block('''\
{
  _id:                ObjectId,
  name:               String,
  email:              String,      // unique index
  phone:              String,      // unique index
  password_hash:      String,      // bcrypt, NEVER returned to client
  preferred_language: "en"|"hi"|"ta"|"te"|"kn"|"ml"|"bn"|"mr"|"gu"|"pa",
  push_token:         String|null, // Expo push token, updated on login
  created_at:         Date,
}
''')

h2('10.2 complaints collection (Nivedan)')
code_block('''\
{
  _id:               ObjectId,
  user_id:           ObjectId,
  category:          String,        // "Labor Issues"
  subcategory:       String,        // "salary_not_paid"
  title:             String,
  status:            "pending"|"filed"|"acknowledged"|"under_review"
                     |"next_step"|"resolved"|"failed",
  form_data:         Object,        // collected fields (grows as agent extracts)
  agent_state:       "CHAT"|"SHOW_BLANK_FORM"|"COLLECT_FIELDS"
                     |"COLLECT_DOCS"|"PREVIEW"|"SUBMITTED"|"REJECTED",
  agent_history:     Array,         // [{role, content}] last 40 messages
  signature_b64:     String|null,   // base64 handwritten signature image
  supporting_docs:   [{             // uploaded supporting documents
    filename:   String,
    file_b64:   String,
    mime_type:  String,
  }],
  portal_ref_id:     String|null,
  rejection_reason:  String|null,
  resubmission_count: Number,
  current_step_label: String,
  timeline:          [{ event, timestamp, detail }],
  documents:         Array,
  created_at:        Date,
  updated_at:        Date,
}
''')

h2('10.3 notifications collection (Nivedan)')
code_block('''\
{
  _id:          ObjectId,
  user_id:      ObjectId,
  complaint_id: String,
  type:         "status_update"|"rejection"|"resolution",
  title:        String,
  body:         String,
  read:         Boolean,        // default false
  created_at:   Date,
}
''')

h2('10.4 portal_submissions collection (Nivedan_portal)')
code_block('''\
{
  _id:            ObjectId,
  complaint_id:   String,       // Nivedan complaint _id
  user_id:        String,
  portal_ref_id:  String,       // "CF-YYYY-NNN"
  category:       String,
  subcategory:    String,
  form_data:      Object,
  pdf_base64:     String,
  status:         String,       // tracks portal-side status
  created_at:     Date,
  updated_at:     Date,
}
''')

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 11. API REFERENCE
# ═══════════════════════════════════════════════════════════════════════════════

h1('11. API Reference')
body('All authenticated endpoints require: Authorization: Bearer <jwt_token>')

h2('11.1 Auth')
two_col_table([
    ('POST /auth/register',  '{ name, email, password, phone, preferred_language } → { token, user }'),
    ('POST /auth/login',     '{ email, password } → { token, user }'),
    ('GET  /auth/me',        '→ { id, name, email, preferred_language }'),
], header=('Endpoint', 'Body / Response'))

h2('11.2 Agent')
code_block('''\
POST /agent/message    [jwt_required]
Body:    { complaint_id: str, message: str }   # message="" for initial greeting
Returns: {
  reply:          str,           # AI response in user\'s language
  action:         null | "show_pdf" | "show_buttons" | "status_update" |
                  "request_signature" | "request_documents",
  action_data: {
    filename?:      str,         # PDF filename
    pdf_base64?:    str,         # base64-encoded PDF
    label?:         "blank_form" | "filled_form",
    buttons?:       str[],       # for show_buttons action
    status?:        str,         # for status_update
    portal_ref_id?: str,
  },
  thinking_steps: str[],
}
''')

h2('11.3 Complaints')
two_col_table([
    ('POST /complaints/create',                   '{ category, subcategory, form_data } → full doc (use ._id!)'),
    ('GET  /complaints/',                          '→ list of user complaints'),
    ('GET  /complaints/mine',                      '→ { complaints, counts: {pending,filed,resolved,failed} }'),
    ('GET  /complaints/<id>',                      '→ full complaint doc with timeline'),
    ('POST /complaints/<id>/submit_to_portal',    '{ pdf_base64 } → { portal_ref_id }'),
    ('POST /complaints/<id>/resubmit',            '{ pdf_base64, changes_summary } → { portal_ref_id }'),
    ('PATCH /complaints/<id>/status',             '{ status } → updated doc'),
    ('POST /complaints/<id>/upload-doc',          '{ type, file_base64, filename, mime_type } → { ok }'),
], header=('Endpoint', 'Body / Response'))

h2('11.4 Notifications, Forms, Users, Webhooks')
two_col_table([
    ('GET  /notifications/mine',        '→ unread notifications (add ?all=1 for all)'),
    ('POST /notifications/read/<id>',   '→ { ok }'),
    ('POST /notifications/read/all',    '→ { ok }'),
    ('GET  /forms/<form_name>',         '→ blank PDF bytes (proxied from portal)'),
    ('POST /forms/generate',            '{ form_name, form_data } → { pdf_base64 }'),
    ('POST /users/push_token',          '{ token } → { ok }'),
    ('POST /webhooks/portal',           '{ portal_ref_id, status, next_step, reason } → { ok }'),
], header=('Endpoint', 'Description'))

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 12. KEY DESIGN DECISIONS
# ═══════════════════════════════════════════════════════════════════════════════

h1('12. Key Design Decisions')

decisions = [
    (
        'uv-only Python dependency management',
        'uv is ~10x faster than pip and produces reproducible lockfiles. '
        'All Python commands in the project use "uv run" or "uv add". pip is never used.'
    ),
    (
        'No ORM (raw pymongo)',
        'MongoDB\'s schemaless nature plus the project\'s small size made an ORM unnecessary overhead. '
        'Schema dicts in models/ serve as documentation without adding abstraction cost.'
    ),
    (
        'Server-side PDF generation only',
        'fpdf2 PDF generation runs on the Flask backend exclusively. '
        'The client never renders a PDF — it receives base64 and either previews or downloads it. '
        'This keeps the mobile bundle small and ensures consistent output across platforms.'
    ),
    (
        'Module-level ref pattern for polling callbacks',
        'handleSendRef and applyResponseRef are plain refs updated on every render. '
        'Polling intervals hold a ref, so they always call the freshest version of the function '
        'without being listed as effect dependencies (which would restart the intervals on every render).'
    ),
    (
        'Separated uploading vs. thinking state',
        'The chat screen has two separate loading states: thinking (AI processing — shows ThinkingStrip) '
        'and uploading (file upload — shows inline card feedback). '
        'Without this separation, users saw "Thinking..." while browsing their camera roll, '
        'which was confusing.'
    ),
    (
        'EXIF orientation correction',
        'Phone photos embed orientation in EXIF metadata rather than rotating pixels. '
        'fpdf2 ignores EXIF, so signatures appeared sideways. '
        'Pillow\'s ImageOps.exif_transpose() applies the rotation before embedding.'
    ),
    (
        'COLLECT_DOCS as a dedicated agent state',
        'Document upload is a distinct state (not just a chat message) to prevent the LLM '
        'from advancing past it. Server-side guards check that signature_b64 is present '
        '(or user explicitly skipped) before transitioning to PREVIEW.'
    ),
    (
        'Expo push + polling fallback',
        'Expo push notifications may not deliver in demo/Wi-Fi environments. '
        'A 10-second polling fallback ensures users always see status updates during live demos.'
    ),
    (
        'Subcategory-driven agent config',
        'SUBCATEGORY_CONFIG in agent_runner.py drives greeting, language, fields, form name, '
        'and authority. Adding a new complaint category is a single dict entry — '
        'no code changes needed elsewhere.'
    ),
    (
        'JWT same secret everywhere',
        'Mobile, web, and backend all use the same JWT_SECRET and HS256 algorithm. '
        'Tokens are 30-day expiry. The JWT_SECRET check at app startup fails loudly '
        'so the server never starts with an insecure empty secret.'
    ),
]

for title, explanation in decisions:
    h3(title)
    body(explanation)

doc.add_page_break()


# ═══════════════════════════════════════════════════════════════════════════════
# 13. ENVIRONMENT SETUP
# ═══════════════════════════════════════════════════════════════════════════════

h1('13. Environment Setup')

h2('13.1 Backend (.env)')
code_block('''\
MONGO_URI=mongodb://localhost:27017/Nivedan
JWT_SECRET=<strong-random-string-min-32-chars>
SARVAM_API_KEY=<from-sarvam-dashboard>          # bare key, no quotes
MOCK_PORTAL_URL=http://localhost:5001
WEBHOOK_SECRET=<random-string>
FLASK_DEBUG=false
''')

h2('13.2 Mobile (.env)')
code_block('''\
EXPO_PUBLIC_API_URL=http://10.30.29.68:5000     # LAN IP of laptop
# Get current IP: ipconfig → Wi-Fi adapter → IPv4 Address
# If IP changes: update this + npx expo start --clear
''')

h2('13.3 Start Commands')
code_block('''\
# Backend
cd Nivedan/backend && uv run python app.py

# Mock Portal
cd Nivedan/mock_portal && uv run python portal_app.py

# Web Dashboard
cd Nivedan/web && npm run dev

# Mobile (Expo Go)
cd Nivedan/mobile && npx expo start --clear
# Press 'a' = Android emulator, scan QR = physical device

# Mobile (dev build — needed for react-native-pdf)
npx expo run:android
''')

h2('13.4 MongoDB — Clean Reset')
code_block('''\
# In mongosh:
use Nivedan
db.complaints.deleteMany({})
db.notifications.deleteMany({})

use Nivedan_portal
db.portal_submissions.deleteMany({})
''')

# ── Final save ─────────────────────────────────────────────────────────────────
output_path = 'Nivedan_Technical_Documentation.docx'
doc.save(output_path)
print(f"Saved: {output_path}")
